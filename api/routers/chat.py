import asyncio
import json
import logging
import os
import re
import threading
import time
import uuid
from typing import Any

from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse, StreamingResponse
from langchain_core.messages import AIMessage, HumanMessage
from pydantic import BaseModel

from agent.logger import LOG_BASE_DIR, bind_log_context, log_to_component, reset_log_context
from agent.utils import (
    activate_run_context,
    build_run_context,
    extract_visible_text,
    is_private_turn,
    reset_run_context,
    strip_scaffold_tags,
)
from api.dependencies import get_agent, get_connection_manager
from tools.chat_history import delete_session, get_session_list, load_session, save_current_session
from tools.user_profile import get_active_profile, reset_profile, set_active_profile
from tools.watch_conditions import capture_watch_conditions, strip_watch_blocks

router = APIRouter()

# How long the stream may go quiet before it re-sends the current phase with an
# elapsed counter, so a long LLM call never looks like a hung run.
HEARTBEAT_SECONDS = 10

# Global chat state
_active_sessions: dict[str, list[dict[str, str]]] = {}
_thread_costs: dict[str, float] = {}
# Per-thread token totals (input+output). Tokens are always accurate regardless of
# model; cost may be 0 when a slot is unpriced — see agent.cost_tracker.
_thread_tokens: dict[str, int] = {}

# Fire-and-forget post-processing tasks (advice ledger + conversation summary),
# held here for their lifetime. asyncio.create_task() only keeps a weak
# reference to the returned Task — without an explicit strong reference held
# somewhere, the event loop can garbage-collect a still-running task between
# awaits, silently dropping the recommendations/summary it was about to persist.
_background_tasks: set[asyncio.Task] = set()

def _spawn_background_task(coro) -> asyncio.Task:
    """asyncio.create_task() wrapper that keeps a strong reference until done."""
    task = asyncio.create_task(coro)
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)
    return task

class ChatRunContext:
    def __init__(self, thread_id: str):
        self.thread_id = thread_id
        self.cancel_event = threading.Event()

_active_chat_runs: dict[str, ChatRunContext] = {}

def request_cancellation(cancel_event: threading.Event):
    cancel_event.set()

def reset_cancellation():
    pass

class AttachmentModel(BaseModel):
    name: str
    type: str
    data: str  # Base64 string or data URL

class ChatRequest(BaseModel):
    message: str
    thread_id: str | None = None
    deep: bool = False
    ghost: bool = False
    request_id: str | None = None
    attachments: list[AttachmentModel] | None = None
    # Opt in to incremental `{"delta": ...}` frames instead of re-sending the
    # whole answer on every token. Defaults to OFF because the wire format is a
    # published contract: the iOS client (separate repo) documents `text` as
    # "the FULL accumulated answer so far (replace, don't append)" and has no
    # delta field, so flipping this by default would leave it blank until the
    # final frame. Clients that understand deltas ask for them.
    stream_deltas: bool = False

def process_attachments(attachments: list[AttachmentModel] | None) -> tuple[str, list[dict]]:
    """
    Processes a list of chat attachments.
    Returns:
        - extra_prompt_text (str): XML formatted text of parsed text/CSV/PDF files.
        - image_blocks (list): List of langchain image_url blocks.
    """
    import base64
    import io

    if not attachments:
        return "", []

    extra_text = []
    image_blocks = []

    for att in attachments:
        name = att.name
        mime_type = att.type
        data_str = att.data

        # Check if it has the data URI scheme prefix (e.g., "data:text/plain;base64,")
        if "," in data_str:
            header, base64_data = data_str.split(",", 1)
        else:
            base64_data = data_str

        try:
            file_bytes = base64.b64decode(base64_data)
        except Exception as e:
            extra_text.append(f"\n[Error decoding attachment {name}: {e}]")
            continue

        if mime_type.startswith("image/"):
            # Image attachments go directly into the content blocks for multimodal LLMs
            image_blocks.append({
                "type": "image_url",
                "image_url": {"url": data_str}
            })
        elif mime_type == "application/pdf" or name.lower().endswith(".pdf"):
            # PDF parsing with fallback
            try:
                import pypdf
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"

                if text_content.strip():
                    extra_text.append(
                        f"\n[File Attachment: {name}]\n"
                        f"<file_content name=\"{name}\" type=\"pdf\">\n"
                        f"{text_content.strip()}\n"
                        f"</file_content>"
                    )
                else:
                    extra_text.append(f"\n[File Attachment: {name} (PDF) contains no readable text]")
            except ImportError:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error: PDF parsing is disabled because 'pypdf' is not installed. To enable, please run 'pip install pypdf'.\n"
                    f"</file_content>"
                )
            except Exception as e:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error parsing PDF: {e}\n"
                    f"</file_content>"
                )
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or name.lower().endswith(".docx"):
            # DOCX Word parsing
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]

                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            table_texts.append(" | ".join(row_cells))

                text_content = "\n".join(paragraphs_text)
                if table_texts:
                    text_content += "\n\n--- Extracted Tables ---\n" + "\n".join(table_texts)

                if text_content.strip():
                    extra_text.append(
                        f"\n[File Attachment: {name}]\n"
                        f"<file_content name=\"{name}\" type=\"docx\">\n"
                        f"{text_content.strip()}\n"
                        f"</file_content>"
                    )
                else:
                    extra_text.append(f"\n[File Attachment: {name} (Word) contains no readable text]")
            except ImportError:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error: Word parsing is disabled because 'python-docx' is not installed.\n"
                    f"</file_content>"
                )
            except Exception as e:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error parsing Word document: {e}\n"
                    f"</file_content>"
                )
        elif mime_type in {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"} or name.lower().endswith((".xlsx", ".xls")):
            # Excel spreadsheet parsing
            try:
                import pandas as pd
                excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
                sheets_text = []
                for sheet_name in excel_file.sheet_names:
                    df = excel_file.parse(sheet_name)
                    csv_data = df.to_csv(index=False)
                    sheets_text.append(f"--- Sheet: {sheet_name} ---\n{csv_data}")

                text_content = "\n\n".join(sheets_text)
                if text_content.strip():
                    extra_text.append(
                        f"\n[File Attachment: {name}]\n"
                        f"<file_content name=\"{name}\" type=\"excel\">\n"
                        f"{text_content.strip()}\n"
                        f"</file_content>"
                    )
                else:
                    extra_text.append(f"\n[File Attachment: {name} (Excel) contains no data]")
            except ImportError:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error: Excel parsing is disabled because 'pandas' or 'openpyxl' is not installed.\n"
                    f"</file_content>"
                )
            except Exception as e:
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\">\n"
                    f"Error parsing Excel spreadsheet: {e}\n"
                    f"</file_content>"
                )
        else:
            # Assume text/data files (CSV, TXT, JSON, PY, MD, etc.)
            try:
                text_content = file_bytes.decode("utf-8", errors="replace")
                extra_text.append(
                    f"\n[File Attachment: {name}]\n"
                    f"<file_content name=\"{name}\" type=\"{mime_type}\">\n"
                    f"{text_content.strip()}\n"
                    f"</file_content>"
                )
            except Exception as e:
                extra_text.append(f"\n[Error reading attachment {name} as text: {e}]")

    return "".join(extra_text), image_blocks

def set_stream_callbacks(on_token=None, on_status=None):
    """Register global callbacks for real-time status and token updates."""
    from agent.utils import set_stream_callbacks as agent_set_callbacks
    agent_set_callbacks(on_token=on_token, on_status=on_status)

@router.post("/api/chat/stop")
async def chat_stop_endpoint(request: Request):
    """Signal the active worker for a specific chat thread to stop processing."""
    thread_id = None
    try:
        # Try to parse JSON, but handle empty or malformed bodies gracefully
        body = await request.body()
        if body:
            payload = json.loads(body)
            thread_id = payload.get("thread_id")
    except Exception as e:
        log_to_component("server", "Chat", f"Non-critical: Stop request body was not valid JSON: {e}", level=logging.DEBUG)

    if thread_id and thread_id in _active_chat_runs:
        request_cancellation(_active_chat_runs[thread_id].cancel_event)
        return {"status": "cancelled", "thread_id": thread_id}

    # If we have a thread_id but it's not in _active_chat_runs, the run already
    # finished or was never started. Do NOT set the global cancel event — that
    # would kill unrelated background tasks (e.g. news feed).
    if thread_id:
        return {"status": "not_found", "thread_id": thread_id}

    # No thread_id at all — cancel all active chat runs individually
    cancelled = []
    for tid, run_ctx in _active_chat_runs.items():
        request_cancellation(run_ctx.cancel_event)
        cancelled.append(tid)
    return {"status": "cancelled", "cancelled_threads": cancelled}

@router.post("/api/chat")
async def chat_endpoint(request: ChatRequest):
    try:
        agent = get_agent()
    except RuntimeError:
        agent = None

    # Always reset cancellation state at the start of a new run
    reset_cancellation()

    from agent.cost_tracker import reset_session_cost
    reset_session_cost()

    # LLM budget circuit breaker (soft gate). If the persistent, restart-safe
    # rolling budget is exceeded, refuse politely rather than feed a possible
    # runaway. The external watchdog enforces a higher HARD ceiling by killing
    # the process; this is the recoverable in-app stop that trips first.
    try:
        from agent import llm_budget
        _budget_reason = llm_budget.over_soft_budget()
    except Exception:
        _budget_reason = ""
    if _budget_reason:
        log_to_component("chat_runtime", "Budget",
            f"Chat refused: LLM budget exceeded ({_budget_reason})", {
                "thread_id": request.thread_id,
                "request_id": request.request_id,
            }, level=logging.WARNING)
        return JSONResponse(
            {"status": "error", "message": (
                f"AI budget limit reached ({_budget_reason}). New AI requests are "
                "paused as a safety limit; this window resets automatically. "
                "Tune the AIDLC_LLM_MAX_* limits in user_data/.env if needed."
            )},
            status_code=429,
        )

    # Critical Safety: Ensure agent is initialized before processing request
    if agent is None:
        log_to_component("chat_runtime", "Request", "Chat request rejected because agent is not initialized", {
            "thread_id": request.thread_id,
            "request_id": request.request_id,
        }, level=logging.ERROR)
        return JSONResponse({"status": "error", "message": "Agent engine is still starting up... Please wait 5 seconds and try again."}, status_code=530)

    thread_id = request.thread_id
    if not thread_id:
        thread_id = str(uuid.uuid4())
    request_id = request.request_id or str(uuid.uuid4())
    run_id = uuid.uuid4().hex[:12]

    # Track user message for auto-save (Restore from disk if missing in memory)
    if thread_id not in _active_sessions:
        existing = load_session(thread_id)
        if existing:
            _active_sessions[thread_id] = existing["messages"]
            _thread_costs.setdefault(thread_id, existing.get("session_cost_cad", 0.0))
            _thread_tokens.setdefault(thread_id, existing.get("session_tokens", 0))
        else:
            _active_sessions[thread_id] = []
    _thread_costs.setdefault(thread_id, 0.0)

    # Add files metadata to the saved chat log to keep history readable without base64 bloat
    history_content = request.message
    if request.attachments:
        file_names = ", ".join(att.name for att in request.attachments)
        history_content += f"\n\n[Attached: {file_names}]"
    _active_sessions[thread_id].append({"role": "user", "content": history_content})

    # Map 'deep' toggle to response length parameter
    response_length = "Detailed (Deep Analysis)" if request.deep else "Concise (Save $$)"

    config = {
        "configurable": {
            "thread_id": thread_id,
            "profile": get_active_profile(),  # resolved by middleware
            "response_length": response_length
        }
    }

    log_to_component("chat_runtime", "Request", "Accepted chat request", {
        "thread_id": thread_id,
        "request_id": request_id,
        "run_id": run_id,
        "deep": request.deep,
        "ghost": request.ghost,
        "response_length": response_length,
        "message_chars": len(request.message or ""),
        "message_preview": (request.message or "")[:160],
        "attachments_count": len(request.attachments) if request.attachments else 0,
    })

    # Parse and compile attachments
    extra_prompt, image_blocks = process_attachments(request.attachments)

    # Construct multimodal content for HumanMessage
    if not image_blocks and not extra_prompt:
        message_content = request.message
    else:
        text_prompt = request.message
        if extra_prompt:
            text_prompt += "\n" + extra_prompt

        message_content = [{"type": "text", "text": text_prompt}]
        for img in image_blocks:
            message_content.append(img)

    # Initialize state with ghost flag. risk_retry_count is explicitly reset to 0
    # here (not just omitted) because it has no reducer — the checkpointer applies
    # last-value-wins, so leaving it out would let a CRITICAL_FAIL retry consumed on
    # an earlier turn silently exhaust the budget for every later turn in the thread.
    initial_state = {
        "messages": [HumanMessage(content=message_content)],
        "ghost": request.ghost,
        "data_context": {},
        "risk_retry_count": 0,
    }

    queue = asyncio.Queue()
    loop = asyncio.get_running_loop()

    # Track all content that was streamed token-by-token so we can deduplicate
    streamed_buffer = {"val": ""}
    # What the CLIENT actually received. streamed_buffer fills from on_token the
    # moment the worker produces a token, whether or not anything is draining the
    # queue — so it is evidence of production, not of delivery. Deduplicating the
    # post-invoke fallback against it meant that when a stream died mid-run, every
    # message the run went on to produce was suppressed as "already covered" and
    # the turn was lost outright (a compliance-retry revision, its verdict and its
    # warning banner, all dropped). Only text that survived a yield lands here.
    delivered_buffer = {"val": ""}
    # Set when the client's stream is torn down before the run finishes. The worker
    # outlives the connection, so it — not the generator — has to own persistence
    # from that point on.
    client_gone = threading.Event()
    stream_stats = {
        "status_count": 0,
        "token_event_count": 0,
        "token_char_count": 0,
        "first_token_preview": None,
        "last_status": None,
        "queue_enqueue_count": 0,
        "thinking_event_count": 0,
        "thinking_char_count": 0,
    }

    def enqueue_stream_item(item: dict | None, source: str):
        """Queue a streaming payload on the event loop and log failures or key milestones."""
        try:
            queue.put_nowait(item)
            stream_stats["queue_enqueue_count"] += 1

            if item is None:
                log_to_component("chat_runtime", "Queue", "Enqueued stream EOF marker", {
                    "thread_id": thread_id,
                    "request_id": request_id,
                    "run_id": run_id,
                    "source": source,
                    "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                })
            elif "status" in item:
                log_to_component("chat_runtime", "Queue", "Enqueued status update for stream", {
                    "thread_id": thread_id,
                    "request_id": request_id,
                    "run_id": run_id,
                    "source": source,
                    "status": item["status"],
                    "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                })
            elif "full_message" in item:
                log_to_component("chat_runtime", "Queue", "Enqueued full message for stream", {
                    "thread_id": thread_id,
                    "request_id": request_id,
                    "run_id": run_id,
                    "source": source,
                    "visible_chars": len(item["full_message"] or ""),
                    "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                })
            elif "thinking" in item:
                # Logged on the FIRST reasoning chunk only. Whether a provider
                # actually returns its reasoning is configuration-dependent
                # (Gemini needs include_thoughts, see _reasoning_kwargs), so an
                # empty trace panel is otherwise indistinguishable from a broken
                # one — this line is what tells the two apart.
                stream_stats["thinking_event_count"] += 1
                stream_stats["thinking_char_count"] += len(item["thinking"] or "")
                if stream_stats["thinking_event_count"] == 1:
                    log_to_component("chat_runtime", "Queue", "Enqueued first reasoning chunk for stream", {
                        "thread_id": thread_id,
                        "request_id": request_id,
                        "run_id": run_id,
                        "source": source,
                        "preview": (item["thinking"] or "")[:120],
                    })
            elif "token" in item and stream_stats["token_event_count"] == 1:
                log_to_component("chat_runtime", "Queue", "Enqueued first token for stream", {
                    "thread_id": thread_id,
                    "request_id": request_id,
                    "run_id": run_id,
                    "source": source,
                    "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                })
        except Exception as e:
            log_to_component("chat_runtime", "QueueError", "Failed to enqueue stream item", {
                "thread_id": thread_id,
                "request_id": request_id,
                "run_id": run_id,
                "source": source,
                "item_type": (
                    "eof" if item is None else
                    "status" if isinstance(item, dict) and "status" in item else
                    "full_message" if isinstance(item, dict) and "full_message" in item else
                    "token" if isinstance(item, dict) and "token" in item else
                    "unknown"
                ),
                "error": str(e),
            }, level=logging.ERROR)

    is_thinking = [False] # Use a list to simulate a nonlocal-state-like behavior in the closure
    thinking_buffer = [""]

    def on_token(token: str):
        stream_stats["token_event_count"] += 1
        stream_stats["token_char_count"] += len(token or "")

        raw_token = str(token or "")

        # 1. Detection of <thinking> tags and redirection
        # We use a stateful approach to handle tokens that might contain or split the tags.
        combined = thinking_buffer[0] + raw_token

        # Check if we are starting a thinking block
        if not is_thinking[0]:
            if "<thinking>" in combined:
                # Split the token: part before <thinking> goes to main stream, part after goes to thinking
                parts = combined.split("<thinking>", 1)
                pre_thinking = parts[0].replace(thinking_buffer[0], "", 1) if parts[0].startswith(thinking_buffer[0]) else parts[0]

                # Send pre-thinking if any
                if pre_thinking:
                    streamed_buffer["val"] += pre_thinking
                    loop.call_soon_threadsafe(enqueue_stream_item, {"token": pre_thinking}, "on_token")

                is_thinking[0] = True
                thinking_buffer[0] = ""

                # Send the remainder of the combined text (after <thinking>) as thinking tokens
                post_thinking = parts[1]
                if post_thinking:
                    loop.call_soon_threadsafe(enqueue_stream_item, {"thinking": post_thinking}, "on_thinking")
                return

            # Handle partial matches (e.g. "<think" is coming)
            elif "<thinking".startswith(combined) or combined.startswith("<thinking"):
                thinking_buffer[0] = combined
                return # Hold this token until we know if it's <thinking>
            else:
                # Not thinking, just pass through
                thinking_buffer[0] = ""
                streamed_buffer["val"] += raw_token
                loop.call_soon_threadsafe(enqueue_stream_item, {"token": raw_token}, "on_token")

        # If already thinking, check for exit tag </thinking>
        else:
            if "</thinking>" in combined:
                parts = combined.split("</thinking>", 1)
                thinking_content = parts[0]

                # Send the content before </thinking>
                if thinking_content:
                    loop.call_soon_threadsafe(enqueue_stream_item, {"thinking": thinking_content}, "on_thinking")

                is_thinking[0] = False
                thinking_buffer[0] = ""

                # The part AFTER </thinking> goes back to the main stream
                post_thinking = parts[1]
                if post_thinking:
                    streamed_buffer["val"] += post_thinking
                    loop.call_soon_threadsafe(enqueue_stream_item, {"token": post_thinking}, "on_token")
                return

            # Handle partial exit matches (e.g. "</think")
            elif "</thinking>".startswith(combined) or combined.startswith("</thinking"):
                thinking_buffer[0] = combined
                return
            else:
                # Still thinking, just stream it
                thinking_buffer[0] = ""
                loop.call_soon_threadsafe(enqueue_stream_item, {"thinking": raw_token}, "on_thinking")

    stream_stats["token_char_count"] = 0 # Reset for current run

    def on_status(msg: str, degraded: bool = False):
        # Clean current prefix [STATUS]
        clean_msg = str(msg).replace("[STATUS]", "").strip()
        stream_stats["status_count"] += 1
        stream_stats["last_status"] = clean_msg
        log_to_component("chat_runtime", "Status", clean_msg, {
            "thread_id": thread_id,
            "request_id": request_id,
            "run_id": run_id,
            "status_index": stream_stats["status_count"],
        })

        # 1. Broadcast to reasoning status bar via WebSockets (Live-Link)
        try:
            loop.call_soon_threadsafe(
                lambda: asyncio.create_task(get_connection_manager().broadcast({"type": "status", "message": clean_msg, "degraded": bool(degraded)}))
            )
        except Exception:
            pass

        # 2. Also put in current request queue for the chat stream. `degraded` is
        # an explicit signal from send_status (NOT inferred from emoji in the
        # text) and drives the permanent DEGRADED header in static/js/chat.js.
        loop.call_soon_threadsafe(enqueue_stream_item, {"status": clean_msg, "degraded": bool(degraded)}, "on_status")

    def on_thinking(text: str):
        """Reasoning-trace sink — a model's thought blocks, straight to the UI.

        Distinct from on_token's <thinking>-tag path: those tags are only ever
        produced by wrapping content AFTER a call returns, so nothing reaches
        this stream that way. Providers emit reasoning as typed content blocks
        mid-stream, and extract_stream_text drops them from the answer; this is
        where they land instead.
        """
        loop.call_soon_threadsafe(enqueue_stream_item, {"thinking": str(text or "")}, "on_thinking")

    existing_run = _active_chat_runs.get(thread_id)
    if existing_run is not None:
        request_cancellation(existing_run.cancel_event)
        log_to_component("chat_runtime", "Request", "Cancelled previous active run for thread", {
            "thread_id": thread_id,
            "request_id": request_id,
            "run_id": run_id,
        }, level=logging.WARNING)

    run_context = build_run_context(on_token=on_token, on_status=on_status, on_thinking=on_thinking)
    _active_chat_runs[thread_id] = run_context

    # Capture the request-scoped profile NOW, while we are still inside the
    # request's context (resolved by profile_middleware). The agent runs in a
    # bare threading.Thread below, and ContextVars do NOT propagate into a
    # manually-created thread — nor would they survive the middleware's
    # reset_profile() that fires when the StreamingResponse is returned. Without
    # re-binding it inside the worker, get_active_profile() falls back to
    # "default" and every tool reads the wrong profile's portfolio.
    active_profile = get_active_profile()

    # Runs exactly once per turn, from whichever side reaches the end first: the
    # generator on a clean close, or the worker when the client dropped mid-run.
    _finalized = threading.Event()

    def _finalize_turn(final_text: str, reason: str):
        """Persist the turn: watch conditions, feedback capture, session, post-processing.

        Callable from the event loop (clean close) or from the worker thread
        (client disconnected), so the only loop-affine step — spawning the
        post-processing task — is scheduled rather than created inline.
        """
        if not final_text.strip() or thread_id not in _active_sessions:
            return
        if _finalized.is_set():
            log_to_component("chat_runtime", "Persistence", "Skipped duplicate turn finalize", {
                "reason": reason,
            })
            return
        _finalized.set()

        clean_final = re.sub(r'<thinking>.*?</thinking>', '', final_text, flags=re.DOTALL).strip()
        clean_final = strip_scaffold_tags(clean_final).strip()
        # Harvest the side-channel BEFORE stripping it: this is the only
        # point where the advisor's own trigger levels are still attached
        # to the answer that committed to them. Local I/O only, so it
        # stays inline rather than racing the post-processing task that
        # would otherwise see the already-stripped text.
        capture_watch_conditions(clean_final, source="chat")
        clean_final = strip_watch_blocks(clean_final).strip()
        # Capture the triggering human message BEFORE appending the assistant
        # reply below, so post-processing can tell a screener-lens turn
        # (external_screen/guru_validation/portfolio_audit) from one that
        # actually authorizes a trade call.
        trigger_messages = _active_sessions[thread_id]
        human_query = str(trigger_messages[-1].get("content", "")) if (
            trigger_messages and trigger_messages[-1].get("role") == "user"
        ) else ""
        # 1.5 FEEDBACK CAPTURE. Every turn is written to the per-profile
        # feedback store as an UNRATED interaction, so a thumb click has
        # something to attach to. This is the finalize point on purpose:
        # `clean_final` here is the answer the user actually saw — after
        # <thinking> removal, scaffold stripping and the watch-block
        # strip — so a rated example is an example of the real product,
        # not of the pre-sanitizer draft. It also runs BEFORE the
        # background post-processing task for the same reason the
        # watch-condition harvest does: local I/O, no race.
        # Ghost/@Private turns are skipped entirely — the store keeps
        # verbatim query and answer, which is exactly what the privacy
        # toggle exists to withhold.
        interaction_id = None
        try:
            if is_private_turn(human_query, request.ghost):
                log_to_component("chat_runtime", "Feedback", "Skipped feedback capture (private turn)", {})
            else:
                from tools.feedback import add_interaction
                interaction_id = add_interaction(
                    user_query=human_query,
                    agent_response=clean_final,
                    thread_id=thread_id,
                    source="chat",
                )
                log_to_component("chat_runtime", "Feedback", "Captured interaction for rating", {
                    "interaction_id": interaction_id,
                    "response_chars": len(clean_final),
                })
        except Exception as e:
            # A feedback write must never cost the user their answer.
            log_to_component("chat_runtime", "Feedback", "Feedback capture failed", {
                "error": str(e),
            }, level=logging.ERROR)
        # Read BEFORE this turn's answer joins the session: pushback has
        # to be attributed to the answer it actually followed, and
        # _active_sessions keeps growing under the background task.
        prior_answer = _prior_assistant_answer(_active_sessions.get(thread_id, []))
        _active_sessions[thread_id].append({"role": "assistant", "content": clean_final})
        try:
            from agent.cost_tracker import get_session_cost, get_session_tokens
            run_cost = get_session_cost()
            _thread_costs[thread_id] = _thread_costs.get(thread_id, 0.0) + run_cost
            _thread_tokens[thread_id] = _thread_tokens.get(thread_id, 0) + get_session_tokens()
            save_current_session(
                thread_id, _active_sessions[thread_id],
                session_cost_cad=_thread_costs.get(thread_id, 0.0),
                session_tokens=_thread_tokens.get(thread_id, 0),
            )
            log_to_component("chat_runtime", "Persistence", "Auto-saved assistant response", {
                "assistant_chars": len(clean_final),
                "session_cost_cad": _thread_costs.get(thread_id, 0.0),
                "reason": reason,
            })
            # Trigger async post-processing for advice outcome ledger,
            # conversation summary and the 1.7 observation log. Scheduled onto
            # the server loop because this may be running in the worker thread,
            # which has no loop of its own for create_task to attach to.
            post_coro = _run_chat_post_processing(
                clean_final, thread_id, human_query,
                ghost=request.ghost,
                interaction_id=interaction_id,
                prior_answer=prior_answer,
            )
            try:
                loop.call_soon_threadsafe(lambda c=post_coro: _spawn_background_task(c))
            except RuntimeError as e:
                post_coro.close()
                log_to_component("chat_runtime", "Persistence", "Post-processing not scheduled (loop closed)", {
                    "error": str(e),
                }, level=logging.WARNING)
        except Exception as e:
            log_to_component("chat_runtime", "Persistence", "Auto-save failed", {
                "error": str(e),
            }, level=logging.ERROR)

    def run_agent():
        prof_token = set_active_profile(active_profile)
        context_token = activate_run_context(run_context)
        log_context_token = bind_log_context(
            thread_id=thread_id,
            request_id=request_id,
            run_id=run_id,
            deep=request.deep,
            response_length=response_length,
        )
        try:
            inputs = initial_state
            log_to_component("chat_runtime", "Invoke", "Invoking agent graph", {
                "message_chars": len(request.message or ""),
                "configurable": config.get("configurable", {}),
            })
            final_state = agent.invoke(inputs, config=config)

            # After invoke, find ALL AI messages after the last human message
            messages = final_state.get("messages", [])
            log_to_component("chat_runtime", "Invoke", "Agent invocation completed", {
                "message_count": len(messages),
            })
            # Deduplicate against what the client RECEIVED, not what this worker
            # produced. The two only diverge once the stream is dead — and there,
            # comparing against production suppressed every remaining message as a
            # duplicate of text nobody ever saw. On a live stream keep using the
            # produced buffer: the generator can legitimately lag a few tokens
            # behind the queue, and a false non-duplicate renders the answer twice.
            streamed_visible_text = extract_visible_text(
                delivered_buffer["val"] if client_gone.is_set() else streamed_buffer["val"],
                strip_node_prefix=True,
            )
            # Visible messages the client did NOT get — assembled the same way the
            # stream assembles them, so a disconnected turn persists as the answer
            # it would have rendered.
            undelivered_parts: list[str] = []
            visible_output_emitted = False
            last_human_idx = -1
            for i, msg in enumerate(messages):
                if isinstance(msg, HumanMessage):
                    last_human_idx = i

            if last_human_idx != -1:
                for msg in messages[last_human_idx + 1:]:
                    if not isinstance(msg, AIMessage) or not msg.content:
                        continue

                    content = msg.content
                    # Handle Bedrock list content
                    if isinstance(content, list):
                        content = "\n".join([item.get("text", "") if isinstance(item, dict) else str(item) for item in content])

                    # Strip [NodeName]: prefix
                    clean_content = re.sub(r'^\[.*?\]:?\s*', '', content)
                    if not clean_content.strip():
                        log_to_component("chat_runtime", "Emit", "Skipped empty post-invoke AI message", {
                            "message_name": getattr(msg, "name", None),
                        })
                        continue

                    # Deduplicate: Check if this message was streamed live
                    # If it was streamed, it will exist in the async token buffer.
                    # Normalize whitespace and strip internal node markers for robust matching
                    def _normalize_for_dedup(s):
                        # Strip all [NodeName]: prefixes anywhere in the text
                        s = re.sub(r'\[.*?\]:?\s*', '', s)
                        # Strip common block markers used for routing
                        s = re.sub(r'---', '', s)
                        # Normalize whitespace
                        return re.sub(r'\s+', ' ', s).strip()

                    norm_streamed = _normalize_for_dedup(streamed_visible_text)
                    norm_content = _normalize_for_dedup(clean_content)

                    # Check multiple snippet lengths for robustness
                    is_duplicate = False
                    if norm_content:
                        # 1. Exact Substring Check (First 80-200 chars)
                        for snip_len in (200, 120, 80):
                            check_snippet = norm_content[:snip_len]
                            if check_snippet and check_snippet in norm_streamed:
                                is_duplicate = True
                                break

                        # 2. Suffix Match (Last 80-200 chars)
                        if not is_duplicate:
                            for snip_len in (200, 120, 80):
                                check_snippet = norm_content[-snip_len:]
                                if check_snippet and check_snippet in norm_streamed:
                                    is_duplicate = True
                                    break

                        # 3. Fuzzy Word Overlap Check (>50% words exist in stream)
                        if not is_duplicate and norm_streamed:
                            content_words = set(norm_content.split())
                            streamed_words = set(norm_streamed.split())
                            if content_words:
                                overlap = len(content_words & streamed_words) / len(content_words)
                                if overlap > 0.5:
                                    is_duplicate = True

                    if is_duplicate:
                        log_to_component("chat_runtime", "Emit", "Skipped post-invoke AI message because visible stream already covered it", {
                            "message_name": getattr(msg, "name", None),
                            "snippet": norm_content[:80],
                        })
                        continue

                    log_to_component("chat_runtime", "Emit", "Queued post-invoke AI message", {
                        "message_name": getattr(msg, "name", None),
                        "visible_chars": len(clean_content),
                        "preview": clean_content[:160],
                    })
                    loop.call_soon_threadsafe(enqueue_stream_item, {"full_message": clean_content}, "post_invoke")
                    undelivered_parts.append(clean_content)
                    visible_output_emitted = True

            if not visible_output_emitted and not streamed_visible_text and not run_context.cancel_event.is_set():
                log_to_component("chat_runtime", "Emit", "No visible AI output detected; enqueuing fallback render message", {
                    "status_count": stream_stats["status_count"],
                    "token_event_count": stream_stats["token_event_count"],
                    "token_char_count": stream_stats["token_char_count"],
                }, level=logging.WARNING)
                loop.call_soon_threadsafe(
                    enqueue_stream_item,
                    {
                        "full_message": (
                            "I completed the analysis, but the final response could not be rendered cleanly. "
                            "Please retry the request."
                        )
                    },
                    "fallback_emit",
                )

            # The client dropped mid-run, so the generator's own auto-save never
            # ran and everything queued above went into a queue with no reader.
            # This worker holds the only complete copy of the turn — persist it
            # here so a reload recovers the answer instead of losing it. Assembled
            # with the stream's own separator so the saved turn reads exactly as
            # the rendered one would have.
            if client_gone.is_set():
                recovered_parts = [p for p in ([delivered_buffer["val"]] + undelivered_parts) if p.strip()]
                if recovered_parts:
                    log_to_component("chat_runtime", "Persistence", "Client disconnected; persisting turn from worker", {
                        "delivered_chars": len(delivered_buffer["val"]),
                        "undelivered_messages": len(undelivered_parts),
                        "undelivered_chars": sum(len(p) for p in undelivered_parts),
                    }, level=logging.WARNING)
                    _finalize_turn("\n\n---\n\n".join(recovered_parts), reason="client_disconnected")

        except Exception as e:
            import importlib.util
            import traceback

            err_str = traceback.format_exc()
            diag: dict[str, Any] = {}
            try:
                spec = importlib.util.find_spec("agent.nodes.market_analyst")
                if spec and spec.origin:
                    diag["market_analyst_module_path"] = spec.origin
                    diag["market_analyst_mtime"] = os.path.getmtime(spec.origin)
            except Exception as e:
                log_to_component("server", "Error", f"Error fetching diagnostic info: {e}", level=logging.WARNING)
            log_to_component("chat_runtime", "InvokeError", "Agent invocation crashed", {
                "error": str(e),
                "traceback": err_str,
                "process_cwd": os.getcwd(),
                "log_base_dir": LOG_BASE_DIR,
                **diag,
            }, level=logging.ERROR)
            # Terminal, unrecovered failure. Emit a structured signal the frontend
            # renders as a persistent error card + a permanent FAILED header — no
            # raw traceback dumped into the chat (the full traceback is in the logs
            # above). str(e) lets the frontend classify a plain-English cause.
            loop.call_soon_threadsafe(enqueue_stream_item, {"fatal_error": str(e)}, "invoke_exception")
        finally:
            notices = list(run_context.notices)  # snapshot before context reset
            log_to_component("chat_runtime", "Invoke", "Agent worker finished", {
                "cancelled": run_context.cancel_event.is_set(),
                "status_count": stream_stats["status_count"],
                "token_event_count": stream_stats["token_event_count"],
                "token_char_count": stream_stats["token_char_count"],
                "last_status": stream_stats["last_status"],
                "notice_count": len(notices),
            })
            reset_run_context(context_token)
            reset_log_context(log_context_token)
            reset_profile(prof_token)
            if _active_chat_runs.get(thread_id) is run_context:
                _active_chat_runs.pop(thread_id, None)
            if notices:
                loop.call_soon_threadsafe(enqueue_stream_item, {"notices": notices}, "worker_notices")
            loop.call_soon_threadsafe(enqueue_stream_item, None, "worker_finally")  # EOF

    threading.Thread(target=run_agent).start()

    async def event_generator():
        log_context_token = bind_log_context(
            thread_id=thread_id,
            request_id=request_id,
            run_id=run_id,
            deep=request.deep,
            response_length=response_length,
        )
        accumulated = ""
        queue_item_count = 0
        token_payload_count = 0
        full_message_count = 0
        # Delta-streaming state. `sent_text` is the sanitized text the client has
        # already been given, which is what a delta is measured against — it is
        # NOT `accumulated` (raw, pre-sanitizer) or `delivered_buffer` (which
        # tracks raw delivery for the disconnect-recovery dedup).
        send_deltas = bool(getattr(request, "stream_deltas", False))
        sent_text = ""
        # Heartbeat state. A single long LLM call (synthesis at max reasoning
        # effort runs for minutes) emits no status, so without this the status
        # line freezes on whatever came last — measured at >30s of dead air in
        # 44% of runs — and the run looks hung. The heartbeat re-sends the last
        # phase with an elapsed counter so the UI always shows progress.
        stream_started = time.monotonic()
        last_status_text = "Working"
        try:
            log_to_component("chat_runtime", "Stream", "Streaming response opened", {})
            yield json.dumps({"thread_id": thread_id}) + "\n"
            while True:
                try:
                    item = await asyncio.wait_for(queue.get(), timeout=HEARTBEAT_SECONDS)
                except TimeoutError:
                    yield json.dumps({
                        "status": last_status_text,
                        "elapsed": int(time.monotonic() - stream_started),
                        "heartbeat": True,
                        "degraded": False,
                    }) + "\n"
                    continue
                if queue_item_count == 0:
                    log_to_component("chat_runtime", "Stream", "Received first queue item for stream", {
                        "item_type": (
                            "eof" if item is None else
                            "status" if "status" in item else
                            "full_message" if "full_message" in item else
                            "token" if "token" in item else
                            "unknown"
                        ),
                    })
                if item is None:
                    break
                queue_item_count += 1

                if "status" in item:
                    last_status_text = item.get("status") or last_status_text
                    yield json.dumps(item) + "\n"
                    continue

                # Terminal-error signal — forward verbatim for the frontend to
                # render as a persistent error card + permanent FAILED header.
                if "fatal_error" in item:
                    yield json.dumps(item) + "\n"
                    continue

                # Reasoning trace. on_token separates <thinking> content out of
                # the visible stream; forward it so the UI can show the live
                # trace. Without this branch the item fell through to the
                # accumulator below and re-sent the UNCHANGED answer text, so
                # the reasoning was discarded and every thinking token cost a
                # full retransmission of the answer.
                if "thinking" in item:
                    yield json.dumps(item) + "\n"
                    continue

                if "token" in item:
                    accumulated += item["token"]
                    token_payload_count += 1
                elif "full_message" in item:
                    separator = "\n\n---\n\n" if accumulated else ""
                    accumulated += separator + item["full_message"]
                    full_message_count += 1

                # Strip <thinking> tags and any leaked prompt-scaffold tags
                # (e.g. RiskManager's "<output_format strict=\"true\">") before sending
                cleaned = re.sub(r'<thinking>.*?</thinking>', '', accumulated, flags=re.DOTALL)
                cleaned = re.sub(r'<thinking>(?:(?!</thinking>).)*$', '', cleaned, flags=re.DOTALL)
                cleaned = strip_scaffold_tags(cleaned)
                # The watch-conditions side-channel (3.3) is machine-readable, not
                # prose — strip it mid-stream (including a half-arrived block) so
                # the JSON never flashes in the chat on its way to the store.
                cleaned = strip_watch_blocks(cleaned)

                # Re-sending the whole answer on every token is O(n^2) on the
                # wire: a 20,000-char answer cost ~50MB, ~550x the text itself,
                # and it grows with the answer. When the client understands
                # deltas, send only what is new.
                #
                # The sanitizers above can REWRITE history, not just extend it —
                # a <thinking> block that closes, or a watch block that completes,
                # retroactively removes text already sent. So the delta is only
                # valid while `cleaned` still extends what was sent; the moment it
                # stops being a continuation, fall back to one full frame and
                # resync. That keeps the client correct without it having to model
                # the sanitizers.
                if send_deltas:
                    if cleaned.startswith(sent_text):
                        delta = cleaned[len(sent_text):]
                        if delta:
                            yield json.dumps({"delta": delta}) + "\n"
                    else:
                        yield json.dumps({"text": cleaned}) + "\n"
                    sent_text = cleaned
                else:
                    yield json.dumps({"text": cleaned}) + "\n"
                # Only after the yield returns is this text actually on the wire.
                delivered_buffer["val"] = accumulated
        except asyncio.CancelledError:
            # The client is gone. Tell the run: is_cancelled() lets the search
            # fan-out in agent/tool_registry.py bail, and it makes the worker's
            # own "cancelled" log line true instead of reporting a clean finish
            # for a turn nobody received. Graph nodes do NOT poll this today, so
            # the run will keep going to completion — which is precisely why the
            # worker has to persist the result below rather than assume the
            # stream delivered it.
            client_gone.set()
            request_cancellation(run_context.cancel_event)
            log_to_component("chat_runtime", "StreamError", "Streaming response generator cancelled", {
                "queue_item_count": queue_item_count,
                "token_payload_count": token_payload_count,
                "full_message_count": full_message_count,
                "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                "delivered_chars": len(delivered_buffer["val"]),
                "produced_chars": len(streamed_buffer["val"]),
            }, level=logging.WARNING)
            raise
        except GeneratorExit:
            # The third shape a dead client takes, and the easiest one to miss:
            # when the server stops consuming the response, the async generator is
            # closed rather than cancelled, and GeneratorExit is a BaseException —
            # so neither the CancelledError handler above nor the Exception
            # handler below sees it. Left unhandled it means the turn is silently
            # abandoned: no persistence, and a worker that still believes its
            # output was delivered.
            client_gone.set()
            request_cancellation(run_context.cancel_event)
            log_to_component("chat_runtime", "StreamError", "Streaming response generator closed by client", {
                "queue_item_count": queue_item_count,
                "token_payload_count": token_payload_count,
                "full_message_count": full_message_count,
                "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                "delivered_chars": len(delivered_buffer["val"]),
                "produced_chars": len(streamed_buffer["val"]),
            }, level=logging.WARNING)
            raise
        except Exception as e:
            import traceback
            # Same contract as the cancellation path. A client disconnect does not
            # always arrive as CancelledError — Starlette surfaces it as
            # ClientDisconnected from the send side — so both teardowns have to
            # hand the run over to the worker identically, or the shape of the
            # disconnect decides whether the turn survives.
            client_gone.set()
            request_cancellation(run_context.cancel_event)
            log_to_component("chat_runtime", "StreamError", "Streaming response generator crashed", {
                "error": str(e),
                "traceback": traceback.format_exc(),
                "delivered_chars": len(delivered_buffer["val"]),
                "produced_chars": len(streamed_buffer["val"]),
                "queue_item_count": queue_item_count,
                "token_payload_count": token_payload_count,
                "full_message_count": full_message_count,
                "queue_enqueue_count": stream_stats["queue_enqueue_count"],
            }, level=logging.ERROR)
            raise
        else:
            # One reconciling full frame at the end of a delta stream. Costs n
            # bytes once, not n per token, and means a client that dropped or
            # mis-applied a delta still lands on exactly the server's text rather
            # than silently keeping a corrupted answer.
            if send_deltas and sent_text:
                yield json.dumps({"text": sent_text, "final": True}) + "\n"

            visible_accumulated = extract_visible_text(accumulated)
            log_to_component("chat_runtime", "Stream", "Streaming response closed", {
                "queue_item_count": queue_item_count,
                "token_payload_count": token_payload_count,
                "full_message_count": full_message_count,
                "queue_enqueue_count": stream_stats["queue_enqueue_count"],
                "visible_chars": len(visible_accumulated),
                "reasoning_chunks": stream_stats["thinking_event_count"],
                "reasoning_chars": stream_stats["thinking_char_count"],
                "early_end": not bool(visible_accumulated),
            }, level=logging.WARNING if not visible_accumulated and not run_context.cancel_event.is_set() else logging.INFO)

            # AUTO-SAVE: Store the assistant response and persist the session.
            # `accumulated` is the delivered text, so this is the answer the user
            # actually saw. The worker owns this instead when the client dropped.
            _finalize_turn(accumulated, reason="stream_closed")
        finally:
            reset_log_context(log_context_token)

    # `Content-Encoding: identity` opts this response out of GZipMiddleware.
    # Starlette's gzip responder writes each chunk into a GzipFile WITHOUT
    # flushing, so a streaming response is held inside the compressor's window
    # until enough bytes accumulate — tokens would arrive in bursts instead of
    # live. Starlette excludes `text/event-stream` by default but NOT
    # `application/x-ndjson`, which is what this endpoint speaks, so the opt-out
    # has to be explicit. The two other StreamingResponses (the CSV template and
    # the report download) are complete buffers with no latency requirement and
    # are deliberately left compressible.
    return StreamingResponse(
        event_generator(),
        media_type="application/x-ndjson",
        headers={"X-Thread-ID": thread_id, "Content-Encoding": "identity"}
    )


# ============================================================
# CHAT MANAGEMENT ENDPOINTS
# ============================================================

@router.post("/api/memory/extract_thesis")
async def extract_thesis_from_chat(request: dict[str, str]):
    text = request.get("text", "")
    if not text:
        return {"error": "No text provided"}

    from tools.memory import extract_thesis_from_text
    extracted = extract_thesis_from_text(text)
    return extracted

@router.get("/api/chats")
async def list_chats():
    """List all saved chat sessions (for sidebar)."""
    sessions = get_session_list()
    return JSONResponse({"sessions": sessions})

@router.get("/api/chats/{session_id}")
async def get_chat(session_id: str):
    """Load a specific chat session's messages and cost."""
    session_data = load_session(session_id)
    if session_data is None:
        return JSONResponse({"error": "Session not found"}, status_code=404)

    messages = session_data["messages"]
    cost = session_data.get("session_cost_cad", 0.0)

    # Also restore to active sessions so continuing the chat auto-saves
    _active_sessions[session_id] = messages.copy()
    _thread_costs[session_id] = cost
    return JSONResponse({"session_id": session_id, "messages": messages, "session_cost_cad": cost})

@router.delete("/api/chats/{session_id}")
async def remove_chat(session_id: str):
    """Delete a saved chat session."""
    success = delete_session(session_id)
    if session_id in _active_sessions:
        del _active_sessions[session_id]
    _thread_costs.pop(session_id, None)
    return JSONResponse({"success": success})


# ============================================================

def _prior_assistant_answer(session_msgs: list[dict[str, Any]]) -> str | None:
    """The advisor's previous answer in this thread, if there was one.

    Roadmap 1.7 uses this to decide whether a correction is pushback AT the
    advisor: "no, that's wrong" opening a thread is about something outside this
    app, and recording it as a grievance against an answer that was never given
    would fabricate the evidence a rule gets drafted from.

    Call it BEFORE appending the current answer to the session.
    """
    seen_user = False
    for msg in reversed(session_msgs or []):
        role = msg.get("role")
        if role == "user":
            seen_user = True
        elif role == "assistant" and seen_user:
            return str(msg.get("content") or "")
    return None


async def _run_chat_post_processing(
    text: str,
    thread_id: str,
    human_query: str = "",
    *,
    ghost: bool = False,
    interaction_id: str | None = None,
    prior_answer: str | None = None,
):
    """
    Asynchronously extracts recommendations, updates the conversation summary and
    records this turn's behavioural observations (roadmap 1.7).

    This is the POST-turn seam, and 1.7 exists because the memory write path was
    on the PRE-turn one: `process_user_message` fires from the first supervisor
    pass, before any tool has run and before the answer exists, so it judges a
    lone message with no conversation and no outcome. Here the full turn is in
    hand and nothing is on the user's critical path.
    """
    from agent.lenses import SCREENER_ONLY_LENSES, extract_lens
    from agent.utils import safe_print

    # Portfolio read is shared by both consumers below and is blocking I/O
    # (CSV + possibly a live broker call), so it happens once, off the loop.
    # None means unreadable, which is NOT the same as holding nothing — see
    # tools/observations.py::load_holdings_map.
    holdings = None
    private = is_private_turn(human_query, ghost)
    if not private:
        from tools.observations import load_holdings_map
        holdings = await asyncio.to_thread(load_holdings_map)

    # 1. RECOMMENDATION EXTRACTION
    # Screener-lens turns (portfolio_audit/external_screen/guru_validation) never
    # authorize a trade call by contract, and RiskManager independently confirms
    # this on the same turn ("no unauthorized trade recommendation made") — logging
    # them into the ledger anyway would contradict that judgment and skew the
    # hit-rate scorecard with candidate-surfacing picks that were never a real call.
    # Each step below is isolated. They share a turn, not a fate: the extractor
    # builds its LLM client OUTSIDE its own try, so an unconfigured provider used
    # to raise straight out of this coroutine and silently take the conversation
    # summary with it. That is worse now than it was — the observation log is
    # deterministic and needs no model at all, so losing it to someone else's
    # provider outage would be losing exactly the tier that still works.
    try:
        if extract_lens(human_query) in SCREENER_ONLY_LENSES:
            safe_print("Skipping recommendation extraction: screener-only lens")
        else:
            await _extract_and_log_recommendations(
                text,
                holdings=holdings,
                thread_id=thread_id,
                interaction_id=interaction_id,
                record_observations=not private,
            )
    except Exception as e:
        safe_print(f"⚠️ Recommendation post-processing failed: {e}")

    # 2. CONVERSATION SUMMARY (unaffected by the lens gate above)
    try:
        await _update_conversation_summary(thread_id)
    except Exception as e:
        safe_print(f"⚠️ Conversation summary post-processing failed: {e}")

    # 3. OBSERVATION LOG (roadmap 1.7). Deterministic detectors only — no model
    # runs here. The defect being fixed is a one-shot LLM judgment on a single
    # turn with no evidence and no review; adding one at a different seam would
    # reproduce it. Judgment happens later, once, in the gated consolidation pass
    # over accumulated evidence a human then confirms.
    if private:
        safe_print("Skipping observation log: private turn")
    else:
        try:
            from tools.observations import observe_turn
            written = await asyncio.to_thread(
                observe_turn,
                human_query,
                thread_id=thread_id,
                interaction_id=interaction_id,
                prior_answer=prior_answer,
                holdings=holdings,
            )
            if written:
                safe_print(
                    "🔎 Observed: " + ", ".join(w["kind"] for w in written)
                )
        except Exception as e:
            # An observation write must never cost the user their answer.
            safe_print(f"⚠️ Observation log post-processing failed: {e}")


async def _extract_and_log_recommendations(
    text: str,
    *,
    holdings: dict[str, float] | None = None,
    thread_id: str | None = None,
    interaction_id: str | None = None,
    record_observations: bool = False,
):
    import json

    from langchain_core.messages import HumanMessage, SystemMessage

    from agent.memory import _content_to_str
    from agent.utils import get_fast_llm, safe_invoke, safe_print
    from tools.market_data import get_realtime_quote
    from tools.memory import add_recommendation

    llm = get_fast_llm()
    system_prompt = (
        "You are a structured data extractor. Your job is to read the conversation response "
        "and extract any actionable investment advice/recommendations to BUY, SELL, HOLD, TRIM, or ADD any ticker/asset. "
        "Each recommendation must contain: ticker, action (BUY, SELL, HOLD, TRIM, ADD), "
        "confidence_grade (HIGH, MEDIUM, LOW), horizon (Short Term, Medium Term, Long Term, or a time horizon description), "
        "and a concise reason. "
        "Provide output as a valid JSON object matching the schema below. "
        "Do not write any other text, explanations, or markdown. Only return valid JSON.\n\n"
        "Schema:\n"
        "{\n"
        "  \"recommendations\": [\n"
        "    {\n"
        "      \"ticker\": \"AAPL\",\n"
        "      \"action\": \"BUY\",\n"
        "      \"confidence_grade\": \"HIGH\",\n"
        "      \"horizon\": \"Medium Term\",\n"
        "      \"reason\": \"reasons\"\n"
        "    }\n"
        "  ]\n"
        "}"
    )

    try:
        loop = asyncio.get_running_loop()
        response = await loop.run_in_executor(
            None,
            lambda: safe_invoke(llm, [
                SystemMessage(content=system_prompt),
                HumanMessage(content=f"Extract recommendations from this text:\n\n{text}")
            ])
        )
        raw = _content_to_str(response.content).strip()
        if raw.startswith("```"):
            lines = raw.split("\n")
            if lines[0].startswith("```json") or lines[0].startswith("```"):
                raw = "\n".join(lines[1:-1]).strip()

        data = json.loads(raw)
        recommendations = data.get("recommendations", [])

        for rec in recommendations:
            ticker = rec.get("ticker", "").strip().upper()
            action = rec.get("action", "").strip().upper()
            reason = rec.get("reason", "").strip()
            confidence = rec.get("confidence_grade", "").strip().upper()
            horizon = rec.get("horizon", "").strip()

            if ticker and action in ["BUY", "SELL", "HOLD", "TRIM", "ADD"]:
                # Fetch price at advice — off the event loop, same as the LLM call
                # above. get_realtime_quote can make several sequential synchronous
                # HTTP requests (FMP + yfinance fallback) on a cache miss, which
                # would otherwise stall every concurrent SSE stream and WebSocket
                # on this single-worker server for seconds per ticker.
                quote = await loop.run_in_executor(None, get_realtime_quote, ticker)
                price = quote.get("price") if isinstance(quote, dict) else None
                if isinstance(price, str):
                    try:
                        price = float(price.replace("$", "").replace(",", "").strip())
                    except ValueError:
                        price = None

                add_recommendation(
                    ticker=ticker,
                    action=action,
                    reason=reason,
                    price_at_advice=price,
                    confidence_grade=confidence,
                    horizon=horizon
                )
                safe_print(f"Recorded recommendation: {action} {ticker} at {price} ({confidence})")

                # 1.7: anchor the position size AT the call, so follow-through is
                # measurable later. It has to be captured now — portfolio_history
                # stores totals, not positions, so what was held on a past date
                # cannot be reconstructed afterwards. `None` when the portfolio
                # was unreadable this turn: unknown, never zero.
                if record_observations:
                    from tools.observations import record_rec_issued
                    record_rec_issued(
                        ticker,
                        action,
                        # A readable portfolio without this name means zero held;
                        # only an unreadable one is unknown.
                        shares_at_advice=(holdings.get(ticker, 0.0) if holdings is not None else None),
                        thread_id=thread_id,
                        interaction_id=interaction_id,
                    )

    except Exception as e:
        safe_print(f"⚠️ Recommendation post-processing failed: {e}")


async def _update_conversation_summary(thread_id: str):
    try:
        from langchain_core.messages import AIMessage as LangAI
        from langchain_core.messages import HumanMessage as LangHuman

        from agent.memory import summarize_messages
        from agent.utils import safe_print
        from tools.memory import add_conversation_summary

        session_msgs = _active_sessions.get(thread_id, [])
        if session_msgs:
            messages = []
            for msg in session_msgs:
                if msg.get("role") == "user":
                    messages.append(LangHuman(content=msg.get("content", "")))
                elif msg.get("role") == "assistant":
                    messages.append(LangAI(content=msg.get("content", "")))

            loop = asyncio.get_running_loop()
            summary = await loop.run_in_executor(None, lambda: summarize_messages(messages))
            if summary:
                add_conversation_summary(summary, thread_id=thread_id)
                safe_print(f"Auto-saved conversation summary for thread {thread_id}")
    except Exception as e:
        from agent.utils import safe_print
        safe_print(f"⚠️ Conversation summary post-processing failed: {e}")
